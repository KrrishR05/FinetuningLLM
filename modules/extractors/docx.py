from __future__ import annotations

import io
from typing import List, Tuple

from modules.cleaner import full_clean
from modules.extractors._base import BaseExtractor
from modules.models import PageContent, Section, TableData


class DOCXExtractor(BaseExtractor):
    PARAGRAPHS_PER_PAGE = 40

    def extract(self, data: bytes) -> Tuple[List[PageContent], List[str]]:
        from docx import Document

        pages: List[PageContent] = []
        warnings: List[str] = []

        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            warnings.append(f"Failed to open DOCX: {e}")
            return pages, warnings

        paragraphs = [p.text for p in doc.paragraphs]
        if not paragraphs and not doc.tables:
            warnings.append("DOCX file has no paragraphs or tables")
            return pages, warnings

        all_sections: List[Section] = []
        for p in doc.paragraphs:
            text = p.text.strip()
            style_name = p.style.name.lower() if p.style else ""
            if text and ("heading" in style_name or style_name.startswith("h")):
                level = 1
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                all_sections.append(Section(
                    title=text,
                    level=level,
                    text="",
                    page_number=1,
                ))

        all_tables: List[TableData] = []
        for tbl in doc.tables:
            tbl_rows = []
            for row in tbl.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                tbl_rows.append(row_cells)

            if tbl_rows:
                headers = tbl_rows[0] if len(tbl_rows) > 1 else []
                data_rows = tbl_rows[1:] if len(tbl_rows) > 1 else tbl_rows
                all_tables.append(TableData(
                    headers=headers,
                    rows=data_rows,
                    page_number=1,
                ))

        if not paragraphs:
            pages.append(PageContent(
                page_number=1,
                text=full_clean("\n".join(t.to_text() for t in all_tables)),
                source_label="Section 1",
                sections=all_sections,
                tables=all_tables,
            ))
            return pages, warnings

        for i in range(0, len(paragraphs), self.PARAGRAPHS_PER_PAGE):
            page_num = (i // self.PARAGRAPHS_PER_PAGE) + 1
            chunk_paras = paragraphs[i : i + self.PARAGRAPHS_PER_PAGE]
            text = full_clean("\n".join(chunk_paras))

            page_sections = [s for s in all_sections if s.page_number == page_num or page_num == 1] if page_num == 1 else []
            page_tables = [t for t in all_tables if t.page_number == page_num or page_num == 1] if page_num == 1 else []

            pages.append(PageContent(
                page_number=page_num,
                text=text,
                source_label=f"Section {page_num}",
                sections=page_sections,
                tables=page_tables,
            ))

        return pages, warnings


def extract_docx(data: bytes) -> Tuple[List[PageContent], List[str]]:
    return DOCXExtractor().extract(data)
